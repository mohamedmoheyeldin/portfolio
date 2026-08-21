#!/usr/bin/env python3
"""Generate local-only tailored PDF and DOCX resumes from verified career facts."""

from __future__ import annotations

import argparse
import json
import re
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer


def load_profile(path: Path) -> dict:
    records = json.loads(path.read_text(encoding="utf-8"))
    return next(record for record in records if record["id"] == "profile")


def safe_slug(value: str) -> str:
    return re.sub(r"[^a-z0-9]+", "-", value.lower()).strip("-")[:60] or "application"


def verified_plan(profile: dict, plan: dict) -> dict:
    known_highlights = {
        highlight for role in profile["experience"] for highlight in role["highlights"]
    }
    known_skills = {
        skill for group in profile["skillGroups"] for skill in group["items"]
    }
    return {
        "company": str(plan.get("company") or "Prospective employer")[:100],
        "role": str(plan.get("role") or profile["headline"])[:120],
        "highlights": [item for item in plan.get("selectedHighlights", []) if item in known_highlights][:8],
        "skills": [item for item in plan.get("selectedSkills", []) if item in known_skills][:16],
    }


def role_for_highlight(profile: dict, highlight: str) -> dict:
    return next(role for role in profile["experience"] if highlight in role["highlights"])


def build_pdf(profile: dict, plan: dict, output: Path) -> None:
    styles = getSampleStyleSheet()
    navy = colors.HexColor("#0B1020")
    blue = colors.HexColor("#1267E8")
    muted = colors.HexColor("#4D5568")
    styles.add(ParagraphStyle(name="Name", parent=styles["Title"], fontName="Helvetica-Bold", fontSize=23, leading=25, textColor=navy, spaceAfter=4))
    styles.add(ParagraphStyle(name="Role", parent=styles["Normal"], fontName="Helvetica-Bold", fontSize=10.5, leading=13, textColor=blue, spaceAfter=7))
    styles.add(ParagraphStyle(name="Section", parent=styles["Heading2"], fontName="Helvetica-Bold", fontSize=9, leading=11, textColor=blue, spaceBefore=9, spaceAfter=4, uppercase=True))
    styles.add(ParagraphStyle(name="BodySmall", parent=styles["Normal"], fontSize=8.4, leading=11.2, textColor=muted, spaceAfter=4))
    styles.add(ParagraphStyle(name="BulletSmall", parent=styles["BodySmall"], leftIndent=10, firstLineIndent=-7, bulletIndent=2, spaceAfter=3))
    document = SimpleDocTemplate(str(output), pagesize=LETTER, rightMargin=.55*inch, leftMargin=.55*inch, topMargin=.48*inch, bottomMargin=.48*inch, title=f"{profile['name']} — {plan['role']}")
    story = [
        Paragraph(profile["name"], styles["Name"]),
        Paragraph(plan["role"], styles["Role"]),
        Paragraph("Reston, VA · mohamedmoheyeldin.jobs@gmail.com · linkedin.com/in/moheyeldin · github.com/mohamedmoheyeldin", styles["BodySmall"]),
        Paragraph("PROFILE", styles["Section"]),
        Paragraph(profile["summary"], styles["BodySmall"]),
        Paragraph("RELEVANT EXPERTISE", styles["Section"]),
        Paragraph(" · ".join(plan["skills"] or profile["competencies"][:7]), styles["BodySmall"]),
        Paragraph("SELECTED EXPERIENCE", styles["Section"]),
    ]
    grouped: dict[str, list[str]] = {}
    for highlight in plan["highlights"]:
        role = role_for_highlight(profile, highlight)
        grouped.setdefault(role["employer"], []).append(highlight)
    if not grouped:
        for role in profile["experience"][:2]:
            grouped[role["employer"]] = role["highlights"][:2]
    for employer, highlights in grouped.items():
        role = next(item for item in profile["experience"] if item["employer"] == employer)
        title = role.get("professionalTitle") or role["title"]
        end = role["end"] or "Present"
        story.append(Paragraph(f"<b>{title}</b> · {employer} · {role['start']}–{end}", styles["BodySmall"]))
        story.extend(Paragraph(f"• {highlight}", styles["BulletSmall"]) for highlight in highlights)
    story.extend([Paragraph("EDUCATION", styles["Section"]), Paragraph("Bachelor's degree, Computer Science · American College of Commerce and Technology", styles["BodySmall"]), Spacer(1, 4), Paragraph(f"Tailored for {plan['company']} · Generated from verified canonical career facts", styles["BodySmall"])])
    document.build(story)


def build_docx(profile: dict, plan: dict, output: Path) -> None:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(.45)
    section.bottom_margin = Inches(.45)
    section.left_margin = Inches(.55)
    section.right_margin = Inches(.55)
    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(9)
    heading = document.add_paragraph()
    run = heading.add_run(profile["name"])
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(11, 16, 32)
    role_run = document.add_paragraph().add_run(plan["role"])
    role_run.bold = True
    role_run.font.color.rgb = RGBColor(18, 103, 232)
    document.add_paragraph("Reston, VA · mohamedmoheyeldin.jobs@gmail.com · linkedin.com/in/moheyeldin · github.com/mohamedmoheyeldin")
    for label, body in [("Profile", profile["summary"]), ("Relevant expertise", " · ".join(plan["skills"] or profile["competencies"][:7]))]:
        document.add_heading(label, level=2)
        document.add_paragraph(body)
    document.add_heading("Selected experience", level=2)
    highlights = plan["highlights"] or [item for role in profile["experience"][:2] for item in role["highlights"][:2]]
    current_employer = None
    for highlight in highlights:
        role = role_for_highlight(profile, highlight)
        if role["employer"] != current_employer:
            current_employer = role["employer"]
            paragraph = document.add_paragraph()
            run = paragraph.add_run(f"{role.get('professionalTitle') or role['title']} · {role['employer']}")
            run.bold = True
        document.add_paragraph(highlight, style="List Bullet")
    document.add_heading("Education", level=2)
    document.add_paragraph("Bachelor's degree, Computer Science · American College of Commerce and Technology")
    document.add_paragraph(f"Tailored for {plan['company']} · Generated from verified canonical career facts")
    document.save(output)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--career", type=Path, required=True)
    parser.add_argument("--plan", type=Path, required=True)
    parser.add_argument("--output", type=Path, required=True)
    args = parser.parse_args()
    profile = load_profile(args.career)
    plan = verified_plan(profile, json.loads(args.plan.read_text(encoding="utf-8")))
    args.output.mkdir(parents=True, exist_ok=True)
    stem = f"mohamed-moheyeldin-{safe_slug(plan['company'])}-{safe_slug(plan['role'])}"
    build_pdf(profile, plan, args.output / f"{stem}.pdf")
    build_docx(profile, plan, args.output / f"{stem}.docx")
    print(json.dumps({"pdf": str(args.output / f"{stem}.pdf"), "docx": str(args.output / f"{stem}.docx")}))


if __name__ == "__main__":
    main()
