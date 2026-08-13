#!/usr/bin/env python3
"""Generate branded one-page and detailed DOCX/PDF resumes from career.json."""

from __future__ import annotations

import argparse
import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib.colors import HexColor
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import KeepTogether, PageBreak, Paragraph, SimpleDocTemplate, Spacer

INK = "0B1020"
MUTED = "4D5568"
BLUE = "1267E8"
VIOLET = "7259FF"
LINE = "DFE3EB"
FONT = "Arial"


def load_profile(source: Path) -> dict:
    records = json.loads(source.read_text(encoding="utf-8"))
    if not records or records[0].get("id") != "profile":
        raise ValueError("career.json must contain the canonical profile record")
    return records[0]


def set_cell_border(paragraph, color=LINE, size="6"):
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "3")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)


def style_run(run, size=9, color=INK, bold=False, italic=False):
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic
    return run


def configure_doc(document: Document, compact: bool):
    section = document.sections[0]
    margin = 0.42 if compact else 0.62
    section.top_margin = Inches(margin)
    section.bottom_margin = Inches(margin)
    section.left_margin = Inches(margin)
    section.right_margin = Inches(margin)
    section.header_distance = Inches(0.25)
    section.footer_distance = Inches(0.25)
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal.font.size = Pt(8.1 if compact else 9.4)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(2.4 if compact else 4.5)
    normal.paragraph_format.line_spacing = 1.04 if compact else 1.12


def add_header(document: Document, profile: dict, compact: bool):
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(1)
    style_run(p.add_run(profile["name"]), 18 if compact else 23, INK, True)
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(3 if compact else 5)
    style_run(p.add_run(profile["headline"]), 9 if compact else 11, BLUE, True)
    contact = f'{profile["location"]}  |  mohamedmoheyeldin.com  |  mohamedmoheyeldin.jobs@gmail.com  |  linkedin.com/in/moheyeldin  |  github.com/mohamedmoheyeldin'
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4 if compact else 8)
    style_run(p.add_run(contact), 6.8 if compact else 8, MUTED)
    set_cell_border(p)


def add_section_heading(document: Document, black: str, accent: str, compact: bool):
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(4 if compact else 10)
    p.paragraph_format.space_after = Pt(2 if compact else 4)
    p.paragraph_format.keep_with_next = True
    style_run(p.add_run(black.upper() + " "), 8.4 if compact else 10.2, INK, True)
    style_run(p.add_run(accent.upper()), 8.4 if compact else 10.2, BLUE, True)
    set_cell_border(p, color=LINE, size="4")


def add_body(document: Document, text: str, compact: bool, italic=False):
    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(2.2 if compact else 5)
    p.paragraph_format.line_spacing = 1.02 if compact else 1.12
    style_run(p.add_run(text), 7.7 if compact else 9.2, MUTED if italic else INK, italic=italic)


def add_bullet(document: Document, text: str, compact: bool):
    p = document.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.16 if compact else 0.22)
    p.paragraph_format.first_line_indent = Inches(-0.12)
    p.paragraph_format.space_after = Pt(1.2 if compact else 3.2)
    p.paragraph_format.line_spacing = 1.0 if compact else 1.08
    style_run(p.add_run(text), 7.4 if compact else 9)


def add_role(document: Document, role: dict, compact: bool, highlights: list[str]):
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(2.8 if compact else 7)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.keep_with_next = True
    title = role.get("professionalTitle") or role["title"]
    style_run(p.add_run(f'{title} | {role["employer"]}'), 8.1 if compact else 10, INK, True)
    date_end = "Present" if role["end"] is None else role["end"]
    style_run(p.add_run(f'  |  {role["start"]} - {date_end}'), 7 if compact else 8.3, MUTED)
    if not compact:
        add_body(document, role["summary"], compact=False, italic=True)
    for item in highlights:
        add_bullet(document, item, compact)


def add_footer(document: Document, label: str):
    for section in document.sections:
        p = section.footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style_run(p.add_run(f"{label} | mohamedmoheyeldin.com"), 7.2, MUTED)


def build_docx(profile: dict, output: Path, compact: bool):
    doc = Document()
    configure_doc(doc, compact)
    add_header(doc, profile, compact)
    add_section_heading(doc, "Professional", "profile", compact)
    add_body(doc, profile["summary"] if compact else "\n\n".join(profile["detailedSummary"]), compact)
    add_section_heading(doc, "Core", "expertise", compact)
    expertise = profile["competencies"][:8] if compact else profile["competencies"]
    add_body(doc, "  |  ".join(expertise), compact)
    if compact:
        add_section_heading(doc, "Selected", "experience", compact)
        limits = [2, 2, 1]
        for role, limit in zip(profile["experience"], limits):
            add_role(doc, role, True, role["highlights"][:limit])
        add_section_heading(doc, "Technical", "toolkit", compact)
        selected = profile["skillGroups"][:4] + profile["skillGroups"][-2:]
        for group in selected:
            add_body(doc, f'{group["label"]}: {", ".join(group["items"][:7])}', True)
        add_section_heading(doc, "Education &", "development", compact)
        education = profile["education"][0]
        add_body(doc, f'{education["credential"]} in {education["field"]} | {education["institution"]} | 2014  |  ' + "  |  ".join(profile["credentials"]), True)
    else:
        add_section_heading(doc, "Technical", "skills", False)
        for group in profile["skillGroups"]:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(3)
            style_run(p.add_run(f'{group["label"]}: '), 9.1, INK, True)
            style_run(p.add_run(", ".join(group["items"])), 9.1, MUTED)
        doc.add_page_break()
        add_section_heading(doc, "Professional", "experience", False)
        for index, role in enumerate(profile["experience"]):
            add_role(doc, role, False, role["highlights"])
            if index == 0:
                doc.add_page_break()
        doc.add_page_break()
        add_section_heading(doc, "AI engineering", "practice", False)
        add_body(doc, profile["aiPractice"]["summary"], False)
        for item in profile["aiPractice"]["highlights"]:
            add_bullet(doc, item, False)
        add_section_heading(doc, "Portfolio", "system", False)
        project = profile["projects"][0]
        add_body(doc, project["description"], False)
        for item in project["highlights"]:
            add_bullet(doc, item, False)
        add_section_heading(doc, "Education &", "credentials", False)
        education = profile["education"][0]
        add_body(doc, f'{education["credential"]} in {education["field"]} | {education["institution"]} | 2014', False)
        for credential in profile["credentials"]:
            add_bullet(doc, credential, False)
    add_footer(doc, "One-page resume" if compact else "Detailed resume")
    doc.core_properties.title = f'{profile["name"]} - {"One-page" if compact else "Detailed"} Resume'
    doc.core_properties.subject = profile["headline"]
    doc.core_properties.author = profile["name"]
    doc.save(output)


def pdf_styles(compact: bool):
    base = getSampleStyleSheet()
    body_size = 7.4 if compact else 9.1
    return {
        "name": ParagraphStyle("Name", parent=base["Normal"], fontName="Helvetica-Bold", fontSize=18 if compact else 23, leading=20 if compact else 26, textColor=HexColor("#" + INK), alignment=TA_CENTER, spaceAfter=2),
        "title": ParagraphStyle("Title", parent=base["Normal"], fontName="Helvetica-Bold", fontSize=9 if compact else 11, leading=11 if compact else 13, textColor=HexColor("#" + BLUE), alignment=TA_CENTER, spaceAfter=3),
        "contact": ParagraphStyle("Contact", parent=base["Normal"], fontName="Helvetica", fontSize=6.7 if compact else 7.8, leading=8 if compact else 10, textColor=HexColor("#" + MUTED), alignment=TA_CENTER, spaceAfter=5),
        "section": ParagraphStyle("Section", parent=base["Normal"], fontName="Helvetica-Bold", fontSize=8.5 if compact else 10.4, leading=10 if compact else 13, textColor=HexColor("#" + INK), spaceBefore=4 if compact else 10, spaceAfter=2 if compact else 4, borderColor=HexColor("#" + LINE), borderWidth=0, borderPadding=1),
        "body": ParagraphStyle("Body", parent=base["Normal"], fontName="Helvetica", fontSize=body_size, leading=8.6 if compact else 11.2, textColor=HexColor("#" + INK), spaceAfter=2 if compact else 4),
        "muted": ParagraphStyle("Muted", parent=base["Normal"], fontName="Helvetica-Oblique", fontSize=body_size, leading=8.6 if compact else 11.2, textColor=HexColor("#" + MUTED), spaceAfter=2 if compact else 4),
        "role": ParagraphStyle("Role", parent=base["Normal"], fontName="Helvetica-Bold", fontSize=8.1 if compact else 10, leading=9.2 if compact else 12, textColor=HexColor("#" + INK), spaceBefore=2 if compact else 7, spaceAfter=1),
        "bullet": ParagraphStyle("Bullet", parent=base["Normal"], fontName="Helvetica", fontSize=7.1 if compact else 8.8, leading=8.2 if compact else 10.7, textColor=HexColor("#" + INK), leftIndent=9 if compact else 13, firstLineIndent=-7, spaceAfter=1 if compact else 3, bulletIndent=1),
    }


def section_pdf(story, styles, black, accent):
    story.append(Paragraph(f'{black.upper()} <font color="#{BLUE}">{accent.upper()}</font>', styles["section"]))


def role_pdf(story, styles, role, compact, highlights):
    title = role.get("professionalTitle") or role["title"]
    date_end = "Present" if role["end"] is None else role["end"]
    story.append(Paragraph(f'{title} | {role["employer"]} <font color="#{MUTED}" size="7">| {role["start"]} - {date_end}</font>', styles["role"]))
    if not compact:
        story.append(Paragraph(role["summary"], styles["muted"]))
    for item in highlights:
        story.append(Paragraph("&bull; " + item, styles["bullet"]))


def build_pdf(profile: dict, output: Path, compact: bool):
    margin = 0.42 * inch if compact else 0.62 * inch
    document = SimpleDocTemplate(str(output), pagesize=letter, rightMargin=margin, leftMargin=margin, topMargin=margin, bottomMargin=margin, title=f'{profile["name"]} Resume', author=profile["name"])
    styles = pdf_styles(compact)
    story = [
        Paragraph(profile["name"], styles["name"]),
        Paragraph(profile["headline"], styles["title"]),
        Paragraph(f'{profile["location"]} | mohamedmoheyeldin.com | mohamedmoheyeldin.jobs@gmail.com | linkedin.com/in/moheyeldin | github.com/mohamedmoheyeldin', styles["contact"]),
    ]
    section_pdf(story, styles, "Professional", "profile")
    story.append(Paragraph(profile["summary"] if compact else "<br/><br/>".join(profile["detailedSummary"]), styles["body"]))
    section_pdf(story, styles, "Core", "expertise")
    expertise = profile["competencies"][:8] if compact else profile["competencies"]
    story.append(Paragraph(" | ".join(expertise), styles["body"]))
    if compact:
        section_pdf(story, styles, "Selected", "experience")
        for role, limit in zip(profile["experience"], [2, 2, 1]):
            role_pdf(story, styles, role, True, role["highlights"][:limit])
        section_pdf(story, styles, "Technical", "toolkit")
        for group in profile["skillGroups"][:4] + profile["skillGroups"][-2:]:
            story.append(Paragraph(f'<b>{group["label"]}:</b> {", ".join(group["items"][:7])}', styles["body"]))
        section_pdf(story, styles, "Education &", "development")
        education = profile["education"][0]
        story.append(Paragraph(f'<b>{education["credential"]} in {education["field"]}</b> | {education["institution"]} | 2014 | ' + " | ".join(profile["credentials"]), styles["body"]))
    else:
        section_pdf(story, styles, "Technical", "skills")
        for group in profile["skillGroups"]:
            story.append(Paragraph(f'<b>{group["label"]}:</b> {", ".join(group["items"])}', styles["body"]))
        story.append(PageBreak())
        section_pdf(story, styles, "Professional", "experience")
        for index, role in enumerate(profile["experience"]):
            role_pdf(story, styles, role, False, role["highlights"])
            if index == 0:
                story.append(PageBreak())
        story.append(PageBreak())
        section_pdf(story, styles, "AI engineering", "practice")
        story.append(Paragraph(profile["aiPractice"]["summary"], styles["body"]))
        for item in profile["aiPractice"]["highlights"]:
            story.append(Paragraph("&bull; " + item, styles["bullet"]))
        section_pdf(story, styles, "Portfolio", "system")
        project = profile["projects"][0]
        story.append(Paragraph(project["description"], styles["body"]))
        for item in project["highlights"]:
            story.append(Paragraph("&bull; " + item, styles["bullet"]))
        section_pdf(story, styles, "Education &", "credentials")
        education = profile["education"][0]
        story.append(Paragraph(f'<b>{education["credential"]} in {education["field"]}</b> | {education["institution"]} | 2014', styles["body"]))
        for credential in profile["credentials"]:
            story.append(Paragraph("&bull; " + credential, styles["bullet"]))
    document.build(story)


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("source", type=Path)
    parser.add_argument("output", type=Path)
    args = parser.parse_args()
    profile = load_profile(args.source)
    args.output.mkdir(parents=True, exist_ok=True)
    targets = [
        ("mohamed-moheyeldin-resume-one-page", True),
        ("mohamed-moheyeldin-resume-detailed", False),
    ]
    for name, compact in targets:
        build_docx(profile, args.output / f"{name}.docx", compact)
        build_pdf(profile, args.output / f"{name}.pdf", compact)


if __name__ == "__main__":
    main()
